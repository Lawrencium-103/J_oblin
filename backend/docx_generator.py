import re
import hashlib
from pathlib import Path
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


CV_PROFILES = [
    {
        "font": "Calibri", "size": 10, "name_size": 16,
        "section_font": "Calibri", "section_size": 10,
        "color": RGBColor(0x1a, 0x1a, 0x1a),
        "accent": RGBColor(0x1a, 0x3a, 0x5c),
        "margins": (0.75, 0.75, 0.75, 0.75),
        "line_spacing": 1.15, "contact_sep": " | ", "name": "modern-blue",
    },
    {
        "font": "Arial", "size": 10, "name_size": 15,
        "section_font": "Arial", "section_size": 10,
        "color": RGBColor(0x1a, 0x1a, 0x1a),
        "accent": RGBColor(0x2d, 0x50, 0x16),
        "margins": (0.7, 0.7, 0.7, 0.7),
        "line_spacing": 1.1, "contact_sep": " \u2022 ", "name": "classic-green",
    },
    {
        "font": "Georgia", "size": 10, "name_size": 17,
        "section_font": "Georgia", "section_size": 10.5,
        "color": RGBColor(0x1a, 0x1a, 0x1a),
        "accent": RGBColor(0x5c, 0x1a, 0x1a),
        "margins": (0.8, 0.8, 0.8, 0.8),
        "line_spacing": 1.15, "contact_sep": "  |  ", "name": "serif-maroon",
    },
    {
        "font": "Tahoma", "size": 10, "name_size": 14,
        "section_font": "Tahoma", "section_size": 10,
        "color": RGBColor(0x1a, 0x1a, 0x1a),
        "accent": RGBColor(0x1a, 0x1a, 0x5c),
        "margins": (0.65, 0.65, 0.65, 0.65),
        "line_spacing": 1.1, "contact_sep": " \u2014 ", "name": "compact-navy",
    },
    {
        "font": "Garamond", "size": 11, "name_size": 18,
        "section_font": "Garamond", "section_size": 11,
        "color": RGBColor(0x1a, 0x1a, 0x1a),
        "accent": RGBColor(0x5c, 0x3a, 0x1a),
        "margins": (0.8, 0.8, 0.8, 0.8),
        "line_spacing": 1.2, "contact_sep": " | ", "name": "classic-warm",
    },
    {
        "font": "Verdana", "size": 9.5, "name_size": 15,
        "section_font": "Verdana", "section_size": 9.5,
        "color": RGBColor(0x1a, 0x1a, 0x1a),
        "accent": RGBColor(0x3a, 0x1a, 0x5c),
        "margins": (0.7, 0.7, 0.7, 0.7),
        "line_spacing": 1.15, "contact_sep": " \u2022 ", "name": "wide-purple",
    },
]


def get_cv_profile(seed_str: str = "") -> dict:
    if not seed_str:
        return CV_PROFILES[0]
    idx = int(hashlib.md5(seed_str.encode()).hexdigest(), 16) % len(CV_PROFILES)
    return CV_PROFILES[idx]


def _safe(text: str) -> str:
    if not text:
        return text
    replacements = {
        "\u2014": "--", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2022": "-", "\u2026": "...",
        "\u00a0": " ", "\u2032": "'", "\u2033": '"', "\u25b6": ">",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", errors="replace").decode("latin-1")


SECTIONS_ORDER = ["summary", "skills", "experience", "education", "certifications", "projects", "languages"]


def _dedup_company(title: str, company: str) -> str:
    """If title already contains the company name, don't append it."""
    if company and company.lower() in title.lower():
        return title
    return f"{title}, {company}" if title and company else (title or company or "")


def _add_right_tab_stop(paragraph, position=Inches(7.0)):
    """Add a right-aligned tab stop for dates."""
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(position, alignment=2)  # 2 = RIGHT


def _is_grouped_skills(skills: list) -> bool:
    return bool(skills) and isinstance(skills[0], dict) and "domain" in skills[0]


def _render_skills_docx(doc, skills: list, max_groups: int = 6, profile: dict = None):
    """Render skills in DOCX — handles both flat strings and grouped dicts."""
    if not skills:
        return
    _add_section_title_docx(doc, "Core Competencies", profile)
    if _is_grouped_skills(skills):
        for group in skills[:max_groups]:
            domain = group.get("domain", group.get("name", ""))
            items = group.get("items", [])
            desc = group.get("description", "")
            line = f"{domain}: {', '.join(items)}" if items else domain
            if desc:
                line += f" — {desc}"
            bp = doc.add_paragraph(line, style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.left_indent = Inches(0.25)
            for run in bp.runs:
                run.font.size = Pt(9)
    else:
        for s in skills[:8]:
            bp = doc.add_paragraph(s, style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.left_indent = Inches(0.25)
            for run in bp.runs:
                run.font.size = Pt(9)


def _render_skills_pdf(pdf, skills: list, max_groups: int = 6, profile: dict = None):
    """Render skills in PDF — handles both flat strings and grouped dicts."""
    if not skills:
        return
    pdf._section_title("Core Competencies")
    if _is_grouped_skills(skills):
        for group in skills[:max_groups]:
            domain = group.get("domain", group.get("name", ""))
            items = group.get("items", [])
            desc = group.get("description", "")
            line = f"{domain}: {', '.join(items)}" if items else domain
            if desc:
                line += f" — {desc}"
            pdf._bullet(line)
    else:
        for s in skills[:8]:
            pdf._bullet(s)
    pdf.ln(1)


def _render_skills_text(skills: list, max_groups: int = 6) -> list:
    """Render skills as text lines for preview — handles both formats."""
    lines = []
    if not skills:
        return lines
    lines.append("SKILLS")
    if _is_grouped_skills(skills):
        for group in skills[:max_groups]:
            domain = group.get("domain", group.get("name", ""))
            items = group.get("items", [])
            desc = group.get("description", "")
            line = f"{domain}: {', '.join(items)}" if items else domain
            if desc:
                line += f" — {desc}"
            lines.append("  * " + line)
    else:
        for s in skills[:8]:
            lines.append(f"  * {s}")
    lines.append("")
    return lines


def _add_section_title_docx(doc, title: str, profile: dict = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(profile.get("section_size", 10)) if profile else Pt(10)
    col = profile.get("accent", RGBColor(0x1a, 0x1a, 0x1a)) if profile else RGBColor(0x1a, 0x1a, 0x1a)
    run.font.color.rgb = col
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1",
        qn("w:color"): "{:02x}{:02x}{:02x}".format(col.red, col.green, col.blue),
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_referees_docx(doc, referees: list = None, profile: dict = None):
    _add_section_title_docx(doc, "Referees", profile)
    if referees and any(r.get("name", "") for r in referees):
        for r in referees[:3]:
            parts = [r.get("name", ""), r.get("title", ""), r.get("company", ""), r.get("email", ""), r.get("phone", "")]
            line = " | ".join(p for p in parts if p)
            if line:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(1)
    else:
        p = doc.add_paragraph("Available upon request.")
    p.paragraph_format.space_after = Pt(3)


def generate_cv_docx(tailored_cv: dict, output_path: str, target_type: str = "local", profile: dict = None) -> str:
    if profile is None:
        profile = CV_PROFILES[0]
    doc = Document()
    m = profile.get("margins", (0.75, 0.75, 0.75, 0.75))
    for section in doc.sections:
        section.top_margin = Inches(m[0])
        section.bottom_margin = Inches(m[1])
        section.left_margin = Inches(m[2])
        section.right_margin = Inches(m[3])

    style = doc.styles["Normal"]
    font = style.font
    font.name = profile["font"]
    font.size = Pt(profile["size"])
    sp = style.paragraph_format
    sp.space_after = Pt(2)
    sp.space_before = Pt(0)
    sp.line_spacing = profile["line_spacing"]

    personal = tailored_cv.get("personal_info", {})
    name = personal.get("name", "")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    website = personal.get("website", "")

    # NAME (largest element — F-pattern anchor)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(name)
    run.bold = True
    run.font.size = Pt(profile["name_size"])
    run.font.color.rgb = profile.get("accent", profile["color"])

    # Contact line
    contact_parts = [p for p in [phone, email, location] if p]
    linkedin_parts = [p for p in [linkedin, website] if p]
    if contact_parts:
        c = doc.add_paragraph()
        c.paragraph_format.space_after = Pt(4)
        run = c.add_run(profile["contact_sep"].join(contact_parts))
        run.font.size = Pt(profile["size"] - 1.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if linkedin_parts:
            run2 = c.add_run("  |  " + profile["contact_sep"].join(linkedin_parts))
            run2.font.size = Pt(profile["size"] - 1.5)
            run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # PROFESSIONAL SUMMARY
    summary = tailored_cv.get("professional_summary", "")
    if summary:
        _add_section_title_docx(doc, "Professional Summary", profile)
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(4)

    # CORE COMPETENCIES
    _render_skills_docx(doc, tailored_cv.get("skills", []), profile=profile)

    # PROFESSIONAL EXPERIENCE
    experience = tailored_cv.get("experience", [])
    if experience:
        _add_section_title_docx(doc, "Professional Experience", profile)
        sz = profile["size"]
        for exp in experience[:3]:
            end = "Present" if exp.get("current") else exp.get("end_date", "")
            date_str = f"{exp.get('start_date', '')} - {end}" if exp.get("start_date") else ""

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            title_company = _dedup_company(exp.get("title", ""), exp.get("company", ""))
            run = p.add_run(title_company)
            run.bold = True
            run.font.size = Pt(sz)
            run.font.name = profile["font"]
            if date_str:
                _add_right_tab_stop(p, Inches(7.0))
                run3 = p.add_run(f"\t{date_str}")
                run3.font.size = Pt(9)
                run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            for ach in exp.get("achievements", [])[:5]:
                bp = doc.add_paragraph(ach, style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.left_indent = Inches(0.25)
                for run_b in bp.runs:
                    run_b.font.size = Pt(9.5)

    # EDUCATION
    education = tailored_cv.get("education", [])
    if education:
        _add_section_title_docx(doc, "Education", profile)
        for edu in education[:3]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            degree_inst = f"{edu.get('degree', '')} \u2014 {edu.get('institution', '')}"
            run = p.add_run(degree_inst)
            run.font.size = Pt(10)
            run.bold = True
            if edu.get("start_date"):
                yr = f"{edu['start_date']} - {edu.get('end_date', '')}"
                _add_right_tab_stop(p, Inches(7.0))
                run2 = p.add_run(f"\t{yr}")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # CERTIFICATIONS
    certifications = tailored_cv.get("certifications", [])
    if certifications:
        _add_section_title_docx(doc, "Certifications", profile)
        certs = [c if isinstance(c, str) else c.get("name", "") for c in certifications[:5]]
        for c in certs:
            bp = doc.add_paragraph(c, style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.left_indent = Inches(0.25)
            for run_c in bp.runs:
                run_c.font.size = Pt(9.5)

    # PROJECTS
    projects = tailored_cv.get("projects", [])
    if projects:
        _add_section_title_docx(doc, "Projects", profile)
        for proj in projects[:4]:
            tech = f" [{', '.join(proj.get('technologies', []))}]" if proj.get("technologies") else ""
            link = f" ({proj['url']})" if proj.get("url") else ""
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run_pname = p.add_run(f"{proj.get('name', '')}: ")
            run_pname.bold = True
            p.add_run(f"{proj.get('description', '')[:250]}{tech}{link}")

    # VOLUNTEER EXPERIENCE
    volunteer = tailored_cv.get("volunteer_experience", tailored_cv.get("volunteer", []))
    if volunteer:
        _add_section_title_docx(doc, "Volunteer Experience", profile)
        for v in volunteer[:3]:
            if isinstance(v, str):
                bp = doc.add_paragraph(v, style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)
                bp.paragraph_format.left_indent = Inches(0.25)
            else:
                title_v = v.get("title", "")
                org = v.get("company", v.get("organization", ""))
                dates = ""
                if v.get("start_date"):
                    dates = f" ({v['start_date']} - {v.get('end_date', '')})"
                bp = doc.add_paragraph(f"{title_v} \u2014 {org}{dates}", style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)
                bp.paragraph_format.left_indent = Inches(0.25)
                for ach in v.get("achievements", [])[:3]:
                    bp2 = doc.add_paragraph(ach, style="List Bullet")
                    bp2.paragraph_format.space_after = Pt(0)
                    bp2.paragraph_format.left_indent = Inches(0.35)

    # REFEREES (local only)
    if target_type == "local":
        _add_referees_docx(doc, tailored_cv.get("referees"), profile)

    doc.save(output_path)
    return str(output_path)


def generate_cover_docx(letter_text: str, output_path: str, personal_info: dict = None, company: str = "", profile: dict = None) -> str:
    if profile is None:
        profile = CV_PROFILES[0]
    doc = Document()
    m = profile.get("margins", (1.0, 1.0, 1.0, 1.0))
    for section in doc.sections:
        section.top_margin = Inches(m[0])
        section.bottom_margin = Inches(m[1])
        section.left_margin = Inches(m[2])
        section.right_margin = Inches(m[3])

    style = doc.styles["Normal"]
    font = style.font
    font.name = profile["font"]
    font.size = Pt(profile["size"] + 1)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # 1. Header (Sender's details)
    if personal_info:
        name = personal_info.get("name", "")
        email = personal_info.get("email", "")
        phone = personal_info.get("phone", "")
        location = personal_info.get("location", "")
        linkedin = personal_info.get("linkedin", "")
        website = personal_info.get("website", "")

        # Bold name
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        run_name = p_name.add_run(name)
        run_name.bold = True
        run_name.font.size = Pt(profile["name_size"])
        run_name.font.color.rgb = profile.get("accent", profile["color"])

        # Contact line
        contact_parts = [p for p in [phone, email, location] if p]
        linkedin_parts = [p for p in [linkedin, website] if p]
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(8)
        run_contact = p_contact.add_run(profile["contact_sep"].join(contact_parts))
        run_contact.font.size = Pt(profile["size"] - 1)
        run_contact.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if linkedin_parts:
            run_link = p_contact.add_run("  |  " + profile["contact_sep"].join(linkedin_parts))
            run_link.font.size = Pt(profile["size"] - 1)
            run_link.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Horizontal Rule
        p_hr = doc.add_paragraph()
        pPr = p_hr._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "aaaaaa",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        p_hr.paragraph_format.space_after = Pt(12)

    # 2. Date
    from datetime import datetime
    date_str = datetime.now().strftime("%B %d, %Y")
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_after = Pt(12)
    run_date = p_date.add_run(date_str)
    run_date.font.size = Pt(11)

    # 3. Recipient
    if company:
        p_rec = doc.add_paragraph()
        p_rec.paragraph_format.space_after = Pt(12)
        run_rec = p_rec.add_run(f"Hiring Committee\n{company}")
        run_rec.font.size = Pt(11)

    # 4. Salutation
    p_sal = doc.add_paragraph()
    p_sal.paragraph_format.space_after = Pt(12)
    salutation = f"Dear Hiring Team at {company}," if company else "Dear Hiring Manager,"
    run_sal = p_sal.add_run(salutation)
    run_sal.font.size = Pt(11)

    # 5. Body paragraphs
    clean_lines = []
    for para in letter_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.lower().startswith(("dear", "to whom", "hello", "hi ")):
            continue
        if para.lower().startswith(("sincerely", "regards", "best regards", "thank you", "respectfully")):
            continue
        clean_lines.append(para)

    for para_text in clean_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(para_text)
        run.font.size = Pt(11)

    # 6. Sign-off
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_after = Pt(2)
    p_close.paragraph_format.space_before = Pt(12)
    run_close = p_close.add_run("Sincerely,")
    run_close.font.size = Pt(11)

    if personal_info:
        p_sender = doc.add_paragraph()
        p_sender.paragraph_format.space_after = Pt(0)
        run_sender = p_sender.add_run(personal_info.get("name", ""))
        run_sender.bold = True
        run_sender.font.size = Pt(11)

    doc.save(output_path)
    return str(output_path)


def generate_cv_preview_text(tailored_cv: dict) -> str:
    lines = []
    personal = tailored_cv.get("personal_info", {})
    lines.append(personal.get("name", ""))
    contact = " | ".join(p for p in [personal.get("phone"), personal.get("email"), personal.get("location")] if p)
    if contact:
        lines.append(contact)
    lines.append("")

    summary = tailored_cv.get("professional_summary", "")
    if summary:
        lines.append("PROFESSIONAL SUMMARY")
        lines.append(summary)
        lines.append("")

    skills = tailored_cv.get("skills", [])
    if skills:
        lines.extend(_render_skills_text(skills))

    experience = tailored_cv.get("experience", [])
    if experience:
        lines.append("EXPERIENCE")
        for exp in experience[:3]:
            end = "Present" if exp.get("current") else exp.get("end_date", "")
            title_company = _dedup_company(exp.get('title', ''), exp.get('company', ''))
            parts = [title_company]
            if exp.get("start_date") or end:
                parts.append(f"({exp.get('start_date', '')} - {end})")
            lines.append("  ".join(p for p in parts if p))
            for ach in exp.get("achievements", [])[:6]:
                lines.append(f"  - {ach}")
        lines.append("")

    education = tailored_cv.get("education", [])
    if education:
        lines.append("EDUCATION")
        for edu in education[:3]:
            line = f"{edu.get('degree', '')} — {edu.get('institution', '')}"
            if edu.get("start_date"):
                line += f"  ({edu['start_date']} - {edu.get('end_date', '')})"
            lines.append(line)
        lines.append("")

    certifications = tailored_cv.get("certifications", [])
    if certifications:
        lines.append("CERTIFICATIONS")
        certs = [c if isinstance(c, str) else c.get("name", "") for c in certifications[:5]]
        lines.append(", ".join(certs))
        lines.append("")

    projects = tailored_cv.get("projects", [])
    if projects:
        lines.append("PROJECTS")
        for proj in projects[:4]:
            tech = f" [{', '.join(proj.get('technologies', [])[:3])}]" if proj.get("technologies") else ""
            link = f" ({proj['url']})" if proj.get("url") else ""
            lines.append(f"  * {proj.get('name', '')}: {proj.get('description', '')[:200]}{tech}{link}")

    volunteer = tailored_cv.get("volunteer_experience", tailored_cv.get("volunteer", []))
    if volunteer:
        lines.append("")
        lines.append("VOLUNTEER EXPERIENCE")
        for v in volunteer[:3]:
            if isinstance(v, str):
                lines.append(f"  * {v}")
            else:
                lines.append(f"  * {v.get('title', '')} — {v.get('company', v.get('organization', ''))}")

    memberships = tailored_cv.get("professional_memberships", [])
    if memberships:
        lines.append("")
        lines.append("PROFESSIONAL MEMBERSHIPS")
        lines.append(", ".join(m if isinstance(m, str) else m.get("name", "") for m in memberships[:5]))

    return "\n".join(lines).strip()


class _CvPdf(FPDF):
    def __init__(self, profile: dict = None):
        super().__init__()
        self.profile = profile or CV_PROFILES[0]
        self.set_auto_page_break(auto=True, margin=14)
        self.set_margins(14, 14, 14)
        self._accent_rgb = (
            self.profile.get("accent", RGBColor(0x1a, 0x1a, 0x1a))
        )
        self._font_name = self.profile.get("font", "Helvetica")
        self._font_pdf = self._pdf_font_name()

    def _pdf_font_name(self):
        name = self._font_name
        if name in ("Calibri", "Verdana", "Tahoma", "Garamond", "Georgia"):
            return "Helvetica"
        return name

    def _section_title(self, title: str):
        self.set_x(self.l_margin)
        self.set_font(self._font_pdf, "B", 9)
        r, g, b = self._accent_rgb.red, self._accent_rgb.green, self._accent_rgb.blue
        self.set_text_color(r, g, b)
        self.cell(0, 5, _safe(title.upper()), new_x="LMARGIN", new_y="NEXT")
        y = self.get_y()
        self.set_draw_color(r, g, b)
        self.line(self.l_margin, y, self.w - self.r_margin, y)
        self.ln(1.5)

    def _body(self, text: str, size=9):
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "", size)
        self.set_text_color(30, 30, 30)
        self.multi_cell(w=0, h=4, text=_safe(text), new_x="LMARGIN", new_y="NEXT")

    def _bold_line(self, text: str):
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(26, 26, 26)
        self.multi_cell(w=0, h=4.5, text=_safe(text), new_x="LMARGIN", new_y="NEXT")

    def _bullet(self, text: str):
        self.set_font("Helvetica", "", 8.5)
        self.set_text_color(30, 30, 30)
        indent = 4
        self.cell(indent, 4, "- ")
        x = self.get_x()
        self.multi_cell(w=0, h=4, text=_safe(text), new_x="LMARGIN", new_y="NEXT")


def _add_referees_pdf(pdf, referees: list = None):
    pdf._section_title("Referees")
    if referees and any(r.get("name", "") for r in referees):
        for r in referees[:3]:
            parts = [r.get("name", ""), r.get("title", ""), r.get("company", ""), r.get("email", ""), r.get("phone", "")]
            line = " | ".join(p for p in parts if p)
            if line:
                pdf._body(line)
    else:
        pdf._body("Available upon request.")
    pdf = _CvPdf(profile)
    pdf.add_page()
    personal = tailored_cv.get("personal_info", {})
    name = personal.get("name", "")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")

    # Name
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(26, 26, 26)
    pdf.cell(0, 7, name, new_x="LMARGIN", new_y="NEXT")
    # Contact
    contact_parts = [p for p in [phone, email, location] if p]
    if contact_parts:
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(85, 85, 85)
        pdf.cell(0, 4, " | ".join(contact_parts), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    summary = tailored_cv.get("professional_summary", "")
    if summary:
        pdf._section_title("Professional Summary")
        pdf._body(summary)
        pdf.ln(1.5)

    _render_skills_pdf(pdf, tailored_cv.get("skills", []), profile=pdf.profile)

    experience = tailored_cv.get("experience", [])
    if experience:
        pdf._section_title("Professional Experience")
        for exp in experience[:3]:
            end = "Present" if exp.get("current") else exp.get("end_date", "")
            date_str = f"{exp.get('start_date', '')} - {end}" if exp.get("start_date") else ""
            title_line = _dedup_company(exp.get("title", ""), exp.get("company", ""))
            if date_str:
                title_line += f"  ({date_str})"
            pdf._bold_line(title_line)
            for ach in exp.get("achievements", [])[:5]:
                pdf._bullet(ach)
            pdf.ln(1)

    education = tailored_cv.get("education", [])
    if education:
        pdf._section_title("Education")
        for edu in education[:3]:
            line = edu.get("degree", "")
            if edu.get("institution"):
                line += f" -- {edu['institution']}"
            if edu.get("start_date"):
                line += f"  ({edu['start_date']} - {edu.get('end_date', '')})"
            pdf._bold_line(line)

    certifications = tailored_cv.get("certifications", [])
    if certifications and pdf.get_y() < 250:
        pdf._section_title("Certifications")
        for c in certifications[:5]:
            name_c = c if isinstance(c, str) else c.get("name", "")
            pdf._bullet(name_c)

    projects = tailored_cv.get("projects", [])
    if projects and pdf.get_y() < 230:
        pdf._section_title("Projects")
        for proj in projects[:4]:
            tech = f" [{', '.join(proj.get('technologies', [])[:3])}]" if proj.get("technologies") else ""
            link = f" ({proj['url']})" if proj.get("url") else ""
            pdf._body(f"{proj.get('name', '')}: {proj.get('description', '')[:200]}{tech}{link}")

    volunteer = tailored_cv.get("volunteer_experience", tailored_cv.get("volunteer", []))
    if volunteer and pdf.get_y() < 230:
        pdf._section_title("Volunteer Experience")
        for v in volunteer[:3]:
            if isinstance(v, str):
                pdf._bullet(v)
            else:
                title_v = v.get("title", "")
                org = v.get("company", v.get("organization", ""))
                dates = f" ({v['start_date']} - {v.get('end_date', '')})" if v.get("start_date") else ""
                pdf._bullet(f"{title_v} -- {org}{dates}")
                for ach in v.get("achievements", [])[:3]:
                    pdf._bullet(ach)

    if target_type == "local" and pdf.get_y() < 260:
        _add_referees_pdf(pdf, tailored_cv.get("referees"))

    pdf.output(output_path)
    return str(output_path)


def generate_cover_pdf(letter_text: str, output_path: str, personal_info: dict = None, company: str = "", profile: dict = None) -> str:
    pdf = _CvPdf(profile)
    pdf.add_page()

    # 1. Header (Sender's details)
    if personal_info:
        name = personal_info.get("name", "")
        email = personal_info.get("email", "")
        phone = personal_info.get("phone", "")
        location = personal_info.get("location", "")
        linkedin = personal_info.get("linkedin", "")
        website = personal_info.get("website", "")

        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(26, 26, 26)
        pdf.cell(0, 6, _safe(name), new_x="LMARGIN", new_y="NEXT")

        contact_parts = [p for p in [phone, email, location] if p]
        linkedin_parts = [p for p in [linkedin, website] if p]
        contact_line = " | ".join(contact_parts)
        if linkedin_parts:
            contact_line += "  |  " + " | ".join(linkedin_parts)

        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(85, 85, 85)
        pdf.cell(0, 4, _safe(contact_line), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # Horizontal divider
        y = pdf.get_y()
        pdf.set_draw_color(170, 170, 170)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(4)

    # 2. Date
    from datetime import datetime
    date_str = datetime.now().strftime("%B %d, %Y")
    pdf.set_font("Helvetica", "", 10.5)
    pdf.set_text_color(30, 30, 30)
    pdf.cell(0, 5, _safe(date_str), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # 3. Recipient
    if company:
        pdf.set_font("Helvetica", "", 10.5)
        pdf.multi_cell(0, 5, _safe(f"Hiring Committee\n{company}"))
        pdf.ln(3)

    # 4. Salutation
    salutation = f"Dear Hiring Team at {company}," if company else "Dear Hiring Manager,"
    pdf.set_font("Helvetica", "", 10.5)
    pdf.cell(0, 5, _safe(salutation), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # 5. Body
    clean_lines = []
    for para in letter_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.lower().startswith(("dear", "to whom", "hello", "hi ")):
            continue
        if para.lower().startswith(("sincerely", "regards", "best regards", "thank you", "respectfully")):
            continue
        clean_lines.append(para)

    for para_text in clean_lines:
        pdf.multi_cell(0, 5.5, _safe(para_text))
        pdf.ln(4)

    # 6. Closing
    pdf.cell(0, 5.5, "Sincerely,", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    if personal_info:
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.cell(0, 5.5, _safe(personal_info.get("name", "")), new_x="LMARGIN", new_y="NEXT")

    pdf.output(output_path)
    return str(output_path)
